import subprocess
from docx import Document
from docx.shared import Inches
import os
import re
from typing import Optional

# Configure logging
import logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentConverter:
    """Service for converting markdown documents to Word format only."""
    def __init__(self):
        """Initialize the document converter for Word documents only."""
        logger.info("DocumentConverter initialized for Word conversion only")

    def convert_to_word(self, markdown_file_path: str, output_path: Optional[str] = None) -> str:
        # ...existing code...
        pass

    def insert_mermaid_diagram_to_docx(self, mermaid_code: str, docx_path: str, image_path: str = "diagram.png"):
        """
        Render Mermaid diagram to image and insert at the end of a docx file.
        """
        # Save mermaid code to temp file
        mermaid_file = "temp_mermaid.mmd"
        with open(mermaid_file, "w", encoding="utf-8") as f:
            f.write(mermaid_code)

        # Render to PNG using mermaid-cli
        subprocess.run(["mmdc", "-i", mermaid_file, "-o", image_path], check=True)

        # Insert image into docx
        doc = Document(docx_path)
        doc.add_picture(image_path, width=Inches(6))
        doc.save(docx_path)

        # Clean up temp files
        os.remove(mermaid_file)
        os.remove(image_path)
    def __init__(self):
        """Initialize the document converter for Word documents only."""
        logger.info("DocumentConverter initialized for Word conversion only")

    def convert_to_word(self, markdown_file_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert markdown file to Word document.
        
        Args:
            markdown_file_path: Path to the markdown file
            output_path: Optional output path for Word doc. If None, uses same name with .docx extension
            
        Returns:
            Path to the generated Word document
        """
        try:
            # Read markdown content
            with open(markdown_file_path, 'r', encoding='utf-8') as file:
                markdown_content = file.read()
            
            # Determine output path
            if output_path is None:
                base_name = os.path.splitext(markdown_file_path)[0]
                output_path = f"{base_name}.docx"
            
            # Create Word document
            doc = Document()
            
            # Parse markdown and convert to Word elements
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    # Add paragraph break for empty lines
                    doc.add_paragraph()
                    continue
                
                # Handle headers
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    header_text = line.lstrip('# ').strip()
                    
                    if level == 1:
                        doc.add_heading(header_text, level=1)
                    elif level == 2:
                        doc.add_heading(header_text, level=2)
                    elif level == 3:
                        doc.add_heading(header_text, level=3)
                    else:
                        doc.add_heading(header_text, level=4)
                
                # Handle code blocks
                elif line.startswith('```'):
                    # Skip code block markers for now
                    continue
                
                # Handle bullet points
                elif line.startswith('- ') or line.startswith('* '):
                    bullet_text = line[2:].strip()
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    self._add_text_with_links(bullet_para, bullet_text)
                
                # Handle numbered lists
                elif re.match(r'^\d+\.\s', line):
                    numbered_text = re.sub(r'^\d+\.\s', '', line)
                    numbered_para = doc.add_paragraph(style='List Number')
                    self._add_text_with_links(numbered_para, numbered_text)
                
                # Regular paragraphs
                else:
                    # Handle inline code and formatting - preserve links
                    if line:
                        para = doc.add_paragraph()
                        self._add_text_with_links(para, line)
            
            # Save document
            doc.save(output_path)
            logger.info(f"Word document generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error converting to Word: {str(e)}")
            raise
    
    def _add_text_with_links(self, paragraph, text: str):
        """Add text to paragraph while preserving hyperlinks for Word."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # Find all markdown links in the text
        link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
        last_end = 0
        
        for match in re.finditer(link_pattern, text):
            # Add text before the link
            before_link = text[last_end:match.start()]
            if before_link:
                run = paragraph.add_run(self._clean_markdown_formatting_no_links(before_link))
            
            # Add the hyperlink
            link_text = match.group(1)
            link_url = match.group(2)
            
            # Create proper hyperlink in Word
            try:
                # For file:/// links, format them properly for Word
                if link_url.startswith('file:///'):
                    # Word prefers this format
                    formatted_url = link_url
                else:
                    formatted_url = link_url
                
                # Create a proper Word hyperlink with line break for clickability
                hyperlink_run = paragraph.add_run(link_text)
                hyperlink_run.font.color.rgb = None  # Default hyperlink color (blue)
                hyperlink_run.font.underline = True
                
                # Add line break after link to make it clickable
                paragraph.add_run("\n")
                
                # Add URL in parentheses for reference
                url_run = paragraph.add_run(f"({formatted_url})")
                url_run.font.size = url_run.font.size * 0.9 if url_run.font.size else None  # Smaller font
                
                # Add another line break for spacing
                paragraph.add_run("\n")
                
            except Exception as e:
                # Fallback: add link text with URL and line breaks
                paragraph.add_run(f"{link_text}\n{link_url}\n\n")
            
            last_end = match.end()
        
        # Add remaining text after the last link
        remaining_text = text[last_end:]
        if remaining_text:
            run = paragraph.add_run(self._clean_markdown_formatting_no_links(remaining_text))
    
    def _clean_markdown_formatting_no_links(self, text: str) -> str:
        """Remove basic markdown formatting except links for Word conversion."""
        # Remove inline code backticks
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # Remove bold/italic markers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)
        
        # Don't remove links here - handle them separately
        return text

    def _clean_markdown_formatting(self, text: str) -> str:
        """Remove basic markdown formatting for Word conversion."""
        # Remove inline code backticks
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # Remove bold/italic markers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)
        
        # Remove links (keep text only) - LEGACY METHOD
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        return text
    
    def convert_documentation_folder(self, folder_path: str, formats: list = ['docx']) -> dict:
        """
        Convert all markdown files in a folder to Word format only.
        
        Args:
            folder_path: Path to folder containing markdown files
            formats: List of formats to convert to (only 'docx' supported)
            
        Returns:
            Dictionary with conversion results
        """
        results = {
            'converted': [],
            'errors': []
        }
        
        if not os.path.exists(folder_path):
            results['errors'].append(f"Folder not found: {folder_path}")
            return results
        
        # Find all markdown files
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                if file.endswith('.md'):
                    md_path = os.path.join(root, file)
                    
                    try:
                        if 'docx' in formats:
                            docx_path = self.convert_to_word(md_path)
                            results['converted'].append({
                                'source': md_path,
                                'output': docx_path,
                                'format': 'docx'
                            })
                        
                        # Skip PDF conversion - not supported
                        if 'pdf' in formats:
                            results['errors'].append({
                                'file': md_path,
                                'error': 'PDF conversion not available - feature removed'
                            })
                            
                    except Exception as e:
                        results['errors'].append({
                            'file': md_path,
                            'error': str(e)
                        })
        
        return results
